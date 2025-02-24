from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from fpdf import FPDF
import requests
import re
import time
import matplotlib.pyplot as plt
import seaborn as sns
import numpy as np
import pandas as pd
import spacy
import os
import yfinance as yf
from datetime import datetime, timedelta
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"


nlp = spacy.load("en_core_web_sm")

# API Credentials
MISTRAL_API_URL = "https://api.mistral.ai/v1/chat/completions"
MISTRAL_API_KEY = "8fOnJn0tHCSjZzxXYQMiHl0MANAHXLvr"  # Replace with your actual API key

def fetch_company_data(ticker_symbol, period='1y'):
    """
    Fetch financial data for a company using Yahoo Finance API
    
    Args:
        ticker_symbol (str): Company stock ticker symbol (e.g., 'AAPL')
        period (str): Time period for historical data ('1d', '5d', '1mo', '3mo', '6mo', '1y', '2y', '5y', '10y', 'ytd', 'max')
    
    Returns:
        dict: Various financial data for the company
    """
    try:
        # Create a Ticker object
        ticker = yf.Ticker(ticker_symbol)
        
        # Get historical market data
        hist = ticker.history(period=period)
        
        # Get company info and financial data
        info = ticker.info
        
        # Get quarterly financials
        financials = ticker.financials
        
        return {
            "history": hist,
            "info": info,
            "financials": financials 
        }
    except Exception as e:
        print(f"Error fetching company data: {e}")
        return None

def plot_stock_price(data, ticker_symbol):
    """Plot the stock price history"""
    plt.figure(figsize=(12, 6))
    plt.plot(data['history']['Close'])
    plt.title(f'{ticker_symbol} Stock Price Over Time')
    plt.xlabel('Date')
    plt.ylabel('Stock Price (USD)')
    plt.grid(True)
    plt.savefig(f'{ticker_symbol}_stock_price.png')
    plt.close()
    return f'{ticker_symbol}_stock_price.png'

def plot_volume(data, ticker_symbol):
    """Plot the trading volume"""
    plt.figure(figsize=(12, 6))
    plt.bar(data['history'].index, data['history']['Volume'])
    plt.title(f'{ticker_symbol} Trading Volume')
    plt.xlabel('Date')
    plt.ylabel('Volume')
    plt.grid(True, axis='y')
    plt.savefig(f'{ticker_symbol}_volume.png')
    plt.close()
    return f'{ticker_symbol}_volume.png'

def fetch_ai_content(user_prompt, max_tokens=8000):
    """Fetches AI-generated content dynamically with structured formatting."""
    headers = {
        "Authorization": f"Bearer {MISTRAL_API_KEY}",
        "Content-Type": "application/json"
    }
    messages = [{"role": "user", "content": user_prompt}]
    payload = {
        "model": "mistral-medium",
        "messages": messages,
        "max_tokens": max_tokens
    }
    retries = 5
    for attempt in range(retries):
        response = requests.post(MISTRAL_API_URL, json=payload, headers=headers)
        if response.status_code == 200:
            return response.json().get("choices", [{}])[0].get("message", {}).get("content", "No content generated.")
        elif response.status_code == 429:
            print(f"Error 429: Rate limit exceeded. Retrying in {2 ** attempt} seconds...")
            time.sleep(2 ** attempt)
        else:
            print(f"Error {response.status_code}: {response.text}")
            return ""
    return "Failed to fetch content after several attempts."


def generate_long_report(user_prompt, target_words=5000):
    """Generates a long report by making multiple API requests until enough content is collected."""
    full_content = ""
    while len(full_content.split()) < target_words:
        additional_content = fetch_ai_content(user_prompt, max_tokens=4000)
        if additional_content.lower().startswith("error") or additional_content == "No content generated.":
            break
        full_content += "\n\n" + additional_content
        time.sleep(1)  
    return full_content.strip()

def extract_numerical_data(content):
    """Extracts numerical data from AI-generated content using NLP."""
    categories = ["Market Growth", "Future Predictions", "Capital Trends"]
    values = {category: None for category in categories}
    
    doc = nlp(content)
    for sent in doc.sents:
        for category in categories:
            if category.lower() in sent.text.lower():
                numbers = [float(tok.text) for tok in sent if tok.like_num]
                if numbers:
                    values[category] = max(numbers)  
    
    
    print(f"Extracted Data: {values}")

    for category in categories:
        if values[category] is None:
            values[category] = np.random.randint(50, 200)  # Fallback random data
    
    return list(values.keys()), list(values.values())


def generate_graph(topic, content):
    """Generates a graph based on extracted data using Matplotlib."""
    categories, values = extract_numerical_data(content)
    plt.figure(figsize=(6, 4))
    plt.bar(categories, values, color=['blue', 'green', 'orange'])
    plt.xlabel("Aspects")
    plt.ylabel("Values")
    plt.title(f"Analysis of {topic}")
    plt.savefig("graph.png")
    plt.close()
    return "graph.png"

def generate_seaborn_graph(topic):
    """Generates a Seaborn statistical plot for EDA."""
    data = pd.DataFrame({
        "Category": ["Market Growth", "Future Predictions", "Capital Trends", "Adoption Rate", "Investment"],
        "Value": np.random.randint(50, 200, size=5)
    })
    plt.figure(figsize=(6, 4))
    sns.barplot(x="Category", y="Value", hue="Category", data=data, palette="coolwarm", legend=False)
    plt.xticks(rotation=45)
    plt.title(f"Statistical Analysis of {topic}")
    plt.savefig("seaborn_graph.png")
    plt.close()
    return "seaborn_graph.png"

def generate_company_comparison(ticker_symbol, competitors=['MSFT', 'GOOGL', 'AMZN']):
    """Generate a comparison chart between the company and its competitors"""
    data = {}
    metrics = ['marketCap', 'forwardPE', 'profitMargins', 'returnOnEquity']
    labels = ['Market Cap (B)', 'Forward P/E', 'Profit Margin', 'ROE']
    
    all_tickers = [ticker_symbol] + competitors
    
    for ticker in all_tickers:
        try:
            ticker_data = yf.Ticker(ticker).info
            data[ticker] = [
                ticker_data.get('marketCap', 0) / 1e9,  
                ticker_data.get('forwardPE', 0),
                ticker_data.get('profitMargins', 0) * 100,  
                ticker_data.get('returnOnEquity', 0) * 100  
            ]
        except Exception as e:
            print(f"Error getting data for {ticker}: {e}")
            data[ticker] = [0, 0, 0, 0]
    
    
    df = pd.DataFrame(data, index=labels)
    
    
    plt.figure(figsize=(10, 8))
    df.plot(kind='bar', figsize=(10, 6))
    plt.title(f'Comparison of {ticker_symbol} with Competitors')
    plt.ylabel('Value')
    plt.xlabel('Metric')
    plt.grid(axis='y', linestyle='--', alpha=0.7)
    plt.tight_layout()
    plt.savefig(f"{ticker_symbol}_comparison.png")
    plt.close()
    
    return f"{ticker_symbol}_comparison.png"

def generate_docx(content, filename, company_data=None, ticker_symbol=None):
    """Generates a structured DOCX report with professional formatting and embedded graphs."""
    doc = Document()
    doc.add_heading("Generated Report", level=1)
    
    
    if company_data and ticker_symbol:
        doc.add_heading(f"Company Overview: {ticker_symbol}", level=2)
        company_info = company_data['info']
        doc.add_paragraph(f"Company Name: {company_info.get('shortName', 'N/A')}")
        doc.add_paragraph(f"Industry: {company_info.get('industry', 'N/A')}")
        doc.add_paragraph(f"Sector: {company_info.get('sector', 'N/A')}")
        doc.add_paragraph(f"Market Cap: ${company_info.get('marketCap', 0)/1e9:.2f}B")
        doc.add_paragraph(f"PE Ratio: {company_info.get('trailingPE', 'N/A')}")
        
        
        stock_img_path = plot_stock_price(company_data, ticker_symbol)
        doc.add_picture(stock_img_path, width=Pt(400))
        
    
        volume_img_path = plot_volume(company_data, ticker_symbol)
        doc.add_picture(volume_img_path, width=Pt(400))
        
        
        comparison_img_path = generate_company_comparison(ticker_symbol)
        doc.add_picture(comparison_img_path, width=Pt(400))
    
    
    for paragraph in content.split("\n\n"):
        if ":" in paragraph:
            main_topic, sub_content = paragraph.split(":", 1)
            doc.add_paragraph(main_topic.strip(), style='Heading 2')
            doc.add_paragraph(sub_content.strip())
        else:
            doc.add_paragraph(paragraph.strip())
    
    
    if os.path.exists("graph.png"):
        doc.add_picture("graph.png", width=Pt(400))
    if os.path.exists("seaborn_graph.png"):
        doc.add_picture("seaborn_graph.png", width=Pt(400))
    
    docx_path = f"{filename}.docx"
    doc.save(docx_path)
    return docx_path

def generate_pdf(content, filename, company_data=None, ticker_symbol=None):
    """Generates a well-formatted PDF with structured headings, text, and embedded graphs."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", style="B", size=16)
    pdf.cell(200, 10, "Generated Report", ln=True, align='C')
    pdf.ln(10)
    
    
    if company_data and ticker_symbol:
        pdf.set_font("Arial", style="B", size=14)
        pdf.cell(0, 10, f"Company Overview: {ticker_symbol}", ln=True)
        pdf.ln(5)
        
        company_info = company_data['info']
        pdf.set_font("Arial", size=12)
        pdf.cell(0, 8, f"Company Name: {company_info.get('shortName', 'N/A')}", ln=True)
        pdf.cell(0, 8, f"Industry: {company_info.get('industry', 'N/A')}", ln=True)
        pdf.cell(0, 8, f"Sector: {company_info.get('sector', 'N/A')}", ln=True)
        pdf.cell(0, 8, f"Market Cap: ${company_info.get('marketCap', 0)/1e9:.2f}B", ln=True)
        pdf.cell(0, 8, f"PE Ratio: {company_info.get('trailingPE', 'N/A')}", ln=True)
        pdf.ln(5)
        
        # Add stock price graph
        stock_img_path = plot_stock_price(company_data, ticker_symbol)
        pdf.image(stock_img_path, x=20, w=170)
        pdf.ln(5)
        
        # Add volume graph
        volume_img_path = plot_volume(company_data, ticker_symbol)
        pdf.image(volume_img_path, x=20, w=170)
        pdf.ln(5)
        
        # Add competitor comparison
        comparison_img_path = generate_company_comparison(ticker_symbol)
        pdf.image(comparison_img_path, x=20, w=170)
        pdf.ln(5)
    
    # Add the AI generated content
    pdf.set_font("Arial", size=12)
    for paragraph in content.split("\n\n"):
        if ":" in paragraph:
            main_topic, sub_content = paragraph.split(":", 1)
            pdf.set_font("Arial", style="B", size=14)
            pdf.multi_cell(0, 8, main_topic.strip().encode('latin-1', 'replace').decode('latin-1'))
            pdf.ln(3)
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 8, sub_content.strip().encode('latin-1', 'replace').decode('latin-1'))
            pdf.ln(5)
        else:
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 8, paragraph.strip().encode('latin-1', 'replace').decode('latin-1'))
            pdf.ln(5)
    


    if os.path.exists("graph.png"):
        pdf.image("graph.png", x=20, w=170)
    if os.path.exists("seaborn_graph.png"):
        pdf.image("seaborn_graph.png", x=20, w=170)
    
    pdf_path = f"{filename}.pdf"
    pdf.output(pdf_path, 'F')
    return pdf_path

def create_prompt_template(topic, ticker_symbol=None):
    """Creates a detailed prompt template based on the given topic, optionally including company-specific information."""
    if ticker_symbol:
        prompt_template = f"""
Create a comprehensive and detailed financial report on {topic} (Ticker: {ticker_symbol}).

Please ensure the report includes the following sections:
1. **Company Overview**: Provide a summary of {topic}'s business model, key products/services, and market position.
2. **Industry Analysis**: Discuss the competitive landscape and industry trends affecting {ticker_symbol}.
3. **Financial Performance**: Analyze key financial metrics, growth rates, and profitability compared to peers.
4. **Investment Thesis**: Outline the bull and bear cases for investing in {ticker_symbol}.
5. **Risk Assessment**: Identify and evaluate key risks facing the company.
6. **Future Outlook**: Project future performance based on current trends and company strategy.
7. **Conclusion**: Summarize the key investment considerations and provide a recommendation.

Provide data-driven insights and cite relevant financial metrics throughout the report.
"""
    else:
        prompt_template = f"""
Create a comprehensive and detailed report on the topic of {topic}.

Please ensure the report includes the following sections:
1. **Introduction**: Provide an overview of {topic}, including its significance and current relevance.
2. **Background**: Discuss the historical context and key developments in {topic}.
3. **Current Trends**: Analyze the latest trends, data, and statistics in {topic}. Provide relevant graphs and charts.
4. **Challenges**: Identify and elaborate on the primary challenges and obstacles in {topic}.
5. **Opportunities**: Highlight potential opportunities and future prospects in {topic}.
6. **Case Studies**: Include real-world examples and case studies to illustrate key points.
7. **Conclusion**: Summarize the key findings and provide recommendations.

Provide citations and references for any data and statistics used in the report.
"""
    return prompt_template.strip()


def main():
    print("\nWelcome to the Financial Report Generator\n")
    print("Choose report type:")
    print("1. Company Financial Report")
    print("2. General Topic Report")
    report_type = input("Enter your choice (1/2): ").strip()
    
    if report_type == "1":
        ticker_symbol = input("Enter company ticker symbol (e.g., AAPL, MSFT): ").strip().upper()
        company_name = input("Enter company name (or press Enter to use ticker): ").strip()
        if not company_name:
            company_name = ticker_symbol
        
        print(f"\nFetching data for {company_name} ({ticker_symbol})...")
        company_data = fetch_company_data(ticker_symbol)
        
        if not company_data:
            print(f"❌ Could not fetch data for {ticker_symbol}. Falling back to general report.")
            user_topic = company_name
            ticker_symbol = None
            company_data = None
        else:
            user_topic = company_name
            print(f"✅ Successfully fetched data for {company_name}")
    else:
        user_topic = input("Enter the topic for the report: ").strip()
        ticker_symbol = None
        company_data = None
    
    filename = re.sub(r'[^a-zA-Z0-9_]', '', user_topic.replace(' ', '_'))[:50].lower()
    print("\nGenerating a detailed report with statistical analysis. Please wait...")
    
    user_prompt = create_prompt_template(user_topic, ticker_symbol)
    content = generate_long_report(user_prompt, target_words=5000)
    graph_path = generate_graph(user_topic, content)
    seaborn_graph_path = generate_seaborn_graph(user_topic)
    
    while True:
        print("\nChoose an option:")
        print("1. Download PDF")
        print("2. Download DOCX")
        print("3. Exit")
        choice = input("Enter your choice: ")
        if choice == "1":
            pdf_path = generate_pdf(content, filename, company_data, ticker_symbol)
            print(f"✅ PDF saved as {pdf_path}")
        elif choice == "2":
            docx_path = generate_docx(content, filename, company_data, ticker_symbol)
            print(f"✅ DOCX saved as {docx_path}")
        elif choice == "3":
            print("Exiting...")
            break
        else:
            print("❌ Invalid choice. Please try again.")

if __name__ == "__main__":
    main()